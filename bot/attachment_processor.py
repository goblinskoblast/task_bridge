# -*- coding: utf-8 -*-
"""
Attachment Processor
Обработка вложений из email (xlsx, xls, docx, txt)
"""

import logging
import io
from typing import Optional, List, Dict, Any
from email.message import Message as EmailMessage

logger = logging.getLogger(__name__)


def extract_text_from_txt(file_content: bytes) -> Optional[str]:
    """
    Извлекает текст из TXT файла

    Args:
        file_content: Содержимое файла в байтах

    Returns:
        Текст из файла или None
    """
    try:
        # Пробуем разные кодировки
        encodings = ['utf-8', 'windows-1251', 'cp1252', 'iso-8859-1']

        for encoding in encodings:
            try:
                text = file_content.decode(encoding)
                return text.strip()
            except UnicodeDecodeError:
                continue

        # Если ничего не сработало, используем замену ошибочных символов
        return file_content.decode('utf-8', errors='replace').strip()

    except Exception as e:
        logger.error(f"Error extracting text from TXT: {e}")
        return None


def extract_text_from_docx(file_content: bytes) -> Optional[str]:
    """
    Извлекает текст из DOCX файла

    Args:
        file_content: Содержимое файла в байтах

    Returns:
        Текст из документа или None
    """
    try:
        from docx import Document

        # Создаем файл в памяти
        file_stream = io.BytesIO(file_content)
        doc = Document(file_stream)

        # Извлекаем текст из всех параграфов
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

        # Извлекаем текст из таблиц
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    tables_text.append(" | ".join(row_text))

        # Объединяем всё
        all_text = paragraphs + tables_text

        return "\n".join(all_text).strip()

    except ImportError:
        logger.error("python-docx library not installed")
        return None
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
        return None


def extract_text_from_xlsx(file_content: bytes) -> Optional[str]:
    """
    Извлекает текст из XLSX файла

    Args:
        file_content: Содержимое файла в байтах

    Returns:
        Текст из таблицы или None
    """
    try:
        from openpyxl import load_workbook

        # Создаем файл в памяти
        file_stream = io.BytesIO(file_content)
        workbook = load_workbook(file_stream, read_only=True, data_only=True)

        all_text = []

        # Обрабатываем все листы
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]

            all_text.append(f"=== {sheet_name} ===")

            # Читаем строки
            for row in sheet.iter_rows(values_only=True):
                # Фильтруем пустые ячейки
                row_values = [str(cell).strip() for cell in row if cell is not None and str(cell).strip()]

                if row_values:
                    all_text.append(" | ".join(row_values))

        return "\n".join(all_text).strip()

    except ImportError:
        logger.error("openpyxl library not installed")
        return None
    except Exception as e:
        logger.error(f"Error extracting text from XLSX: {e}")
        return None


def extract_text_from_xls(file_content: bytes) -> Optional[str]:
    """
    Извлекает текст из XLS файла (старый формат Excel)

    Args:
        file_content: Содержимое файла в байтах

    Returns:
        Текст из таблицы или None
    """
    try:
        import xlrd

        # Создаем файл в памяти
        file_stream = io.BytesIO(file_content)
        workbook = xlrd.open_workbook(file_contents=file_stream.read())

        all_text = []

        # Обрабатываем все листы
        for sheet in workbook.sheets():
            all_text.append(f"=== {sheet.name} ===")

            # Читаем строки
            for row_idx in range(sheet.nrows):
                row_values = []

                for col_idx in range(sheet.ncols):
                    cell = sheet.cell(row_idx, col_idx)

                    if cell.value:
                        row_values.append(str(cell.value).strip())

                if row_values:
                    all_text.append(" | ".join(row_values))

        return "\n".join(all_text).strip()

    except ImportError:
        logger.error("xlrd library not installed")
        return None
    except Exception as e:
        logger.error(f"Error extracting text from XLS: {e}")
        return None


def process_attachment(filename: str, file_content: bytes) -> Optional[str]:
    """
    Обрабатывает вложение и извлекает текст

    Args:
        filename: Имя файла
        file_content: Содержимое файла в байтах

    Returns:
        Извлеченный текст или None
    """
    filename_lower = filename.lower()

    # Определяем тип файла по расширению
    if filename_lower.endswith('.txt'):
        logger.info(f"Processing TXT file: {filename}")
        return extract_text_from_txt(file_content)

    elif filename_lower.endswith('.docx'):
        logger.info(f"Processing DOCX file: {filename}")
        return extract_text_from_docx(file_content)

    elif filename_lower.endswith('.xlsx'):
        logger.info(f"Processing XLSX file: {filename}")
        return extract_text_from_xlsx(file_content)

    elif filename_lower.endswith('.xls'):
        logger.info(f"Processing XLS file: {filename}")
        return extract_text_from_xls(file_content)

    else:
        logger.warning(f"Unsupported file type: {filename}")
        return None


def extract_attachments_from_email(msg: EmailMessage) -> List[Dict[str, Any]]:
    """
    Извлекает вложения из email сообщения.
    Для поддерживаемых офисных форматов дополнительно извлекает текст,
    но для скачивания сохраняет байты любого attachment.
    """
    attachments = []

    try:
        for part in msg.walk():
            if part.get_content_disposition() != 'attachment':
                continue

            filename = part.get_filename()
            if not filename:
                continue

            from bot.email_handler import decode_mime_header
            filename = decode_mime_header(filename)

            file_content = part.get_payload(decode=True)
            if not file_content:
                continue

            filename_lower = filename.lower()
            supported = filename_lower.endswith(('.txt', '.docx', '.xlsx', '.xls'))
            extracted_text = process_attachment(filename, file_content) if supported else None

            attachments.append({
                'filename': filename,
                'content_type': part.get_content_type(),
                'size': len(file_content),
                'text': extracted_text,
                'file_data': file_content
            })

            logger.info(f"Processed attachment: {filename} ({len(file_content)} bytes)")

    except Exception as e:
        logger.error(f"Error extracting attachments: {e}", exc_info=True)

    return attachments

def format_attachments_text(attachments: List[Dict[str, Any]]) -> str:
    """
    Форматирует текст из вложений для добавления в описание задачи

    Args:
        attachments: Список вложений с извлеченным текстом

    Returns:
        Форматированный текст
    """
    if not attachments:
        return ""

    parts = []

    for attachment in attachments:
        filename = attachment['filename']
        text = attachment.get('text', '')

        if text:
            parts.append(f"\n\n📎 Вложение: {filename}\n{'='*50}\n{text}")
        else:
            parts.append(f"\n\n📎 Вложение: {filename} (не удалось извлечь текст)")

    return "".join(parts)
